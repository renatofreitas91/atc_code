from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]

DOCUMENTS = [
    ("docs/en/User_Guide_Full.md", "docs/docx/ATC_Full_User_Guide_EN.docx", "ATC Full User Guide"),
    ("docs/pt-PT/User_Guide_Full.md", "docs/docx/ATC_Full_User_Guide_PT-PT.docx", "Guia Completo de Utilizador do ATC"),
]


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text


def add_code_block(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.style = document.styles["No Spacing"]


def markdown_to_docx(markdown: str, title: str) -> Document:
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = "Renato Alexandre dos Santos Freitas"

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    document.add_heading(title, level=0)

    in_code = False
    code_lines: list[str] = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(clean_inline(heading.group(2)), level=level)
            continue

        list_item = re.match(r"^\s*[-*]\s+(.*)$", line)
        if list_item:
            document.add_paragraph(clean_inline(list_item.group(1)), style="List Bullet")
            continue

        document.add_paragraph(clean_inline(line))

    if code_lines:
        add_code_block(document, code_lines)

    return document


def build_docx(source: Path, target: Path, title: str) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    markdown = source.read_text(encoding="utf-8")
    document = markdown_to_docx(markdown, title)
    document.save(target)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build ATC DOCX documentation.")
    parser.add_argument("--check", action="store_true", help="Only check that expected DOCX files exist.")
    args = parser.parse_args()

    if args.check:
        missing = [target for _, target, _ in DOCUMENTS if not (ROOT / target).exists()]
        if missing:
            for target in missing:
                print(f"missing: {target}")
            return 1
        print("All documentation DOCX files exist.")
        return 0

    for source_rel, target_rel, title in DOCUMENTS:
        build_docx(ROOT / source_rel, ROOT / target_rel, title)
        print(f"built {target_rel}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
